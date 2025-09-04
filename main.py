from docx import Document

document = Document()
document.add_heading('My Document Title', level=0) # Adds a main title
document.add_paragraph('This is a plain paragraph of text.')

# Adding bold and italic text within a paragraph
p = document.add_paragraph('This paragraph has ')
p.add_run('bold text').bold = True
p.add_run(' and ')
p.add_run('italic text').italic = True

document.save('my_document.docx')
print("Document created successfully!")